from io import BytesIO

import pytest

from ingestion.parser import DOCX_MIME_TYPE, UnsupportedDocumentError, parse_document


def _docx_bytes(build) -> bytes:
    pytest.importorskip("docx")
    from docx import Document as DocxDocument

    document = DocxDocument()
    build(document)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_docx_headings_become_markdown_sections() -> None:
    def build(document) -> None:
        document.add_heading("Chinh sach VPN", level=1)
        document.add_paragraph("Nhan vien phai dung MFA.")
        document.add_heading("Ngoai le", level=2)
        document.add_paragraph("Truong hop khan cap.")

    text, mime_type = parse_document("policy.docx", _docx_bytes(build))

    assert mime_type == DOCX_MIME_TYPE
    assert "# Chinh sach VPN" in text
    assert "## Ngoai le" in text
    assert "Nhan vien phai dung MFA." in text


def test_docx_table_cells_are_extracted() -> None:
    def build(document) -> None:
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Muc"
        table.cell(0, 1).text = "So ngay"
        table.cell(1, 0).text = "Nghi phep nam"
        table.cell(1, 1).text = "15"

    text, _ = parse_document("bang.docx", _docx_bytes(build))

    assert "Muc | So ngay" in text
    assert "Nghi phep nam | 15" in text


def test_docx_preserves_order_of_paragraphs_and_tables() -> None:
    def build(document) -> None:
        document.add_paragraph("Truoc bang")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Trong bang"
        document.add_paragraph("Sau bang")

    text, _ = parse_document("thutu.docx", _docx_bytes(build))

    assert text.index("Truoc bang") < text.index("Trong bang") < text.index("Sau bang")


def test_corrupt_docx_is_rejected() -> None:
    with pytest.raises(UnsupportedDocumentError, match="Unable to parse DOCX"):
        parse_document("corrupt.docx", b"PK\x03\x04 not really a docx")


def test_legacy_doc_is_rejected() -> None:
    with pytest.raises(UnsupportedDocumentError, match="Only .md, .txt, .pdf and .docx"):
        parse_document("legacy.doc", b"\xd0\xcf\x11\xe0")
