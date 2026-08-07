from io import BytesIO
from pathlib import Path

import pytest

from domain.schemas import DocumentStatus
from ingestion.chunker import ChunkingConfig
from ingestion.raptor import RaptorConfig
from ingestion.service import IngestionService
from providers.base import GenerationRequest, GenerationResult
from retrieval.memory_store import MemoryChunkStore
from storage.registry import DocumentRegistry


class StubGenerationProvider:
    name = "stub"

    def generate(self, request: GenerationRequest) -> GenerationResult:
        return GenerationResult(text="tóm tắt cụm", provider="stub", model="stub-model")

    def generate_structured(self, request, response_model):  # pragma: no cover - unused
        raise NotImplementedError


def _docx_bytes(build) -> bytes:
    pytest.importorskip("docx")
    from docx import Document as DocxDocument

    document = DocxDocument()
    build(document)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_reingesting_same_content_is_idempotent(tmp_path: Path) -> None:
    store = MemoryChunkStore()
    service = IngestionService(DocumentRegistry(tmp_path / "registry.json"), store)
    content = b"# Nghi phep\n\nNhan vien duoc nghi 15 ngay."

    first = service.ingest_bytes("policy.md", content)
    second = service.ingest_bytes("policy.md", content)

    assert first.id == second.id
    assert first.version == second.version == 1
    assert len(store.all_chunks) == 1


def test_new_content_replaces_active_document_version(tmp_path: Path) -> None:
    store = MemoryChunkStore()
    service = IngestionService(DocumentRegistry(tmp_path / "registry.json"), store)

    old = service.ingest_bytes("policy.md", b"Old policy")
    new = service.ingest_bytes("policy.md", b"New policy")

    assert new.id == old.id
    assert new.version == 2
    assert {chunk.text for chunk in store.all_chunks} == {"New policy"}


def test_generic_heading_creates_chunk_with_heading_section_and_coordinates(
    tmp_path: Path,
) -> None:
    """Generic (non-legal) headings remain available as chunk sections."""
    def build(document) -> None:
        document.add_heading("Quy trinh su co", level=1)
        document.add_paragraph("Bao cao trong 24 gio.")

    store = MemoryChunkStore()
    service = IngestionService(DocumentRegistry(tmp_path / "registry.json"), store)

    document = service.ingest_bytes("suco.docx", _docx_bytes(build))

    assert document.status is DocumentStatus.READY
    chunks = store.all_chunks
    assert len(chunks) == 1
    assert chunks[0].section == "Quy trinh su co"
    assert chunks[0].coordinates.doc_id == "suco.docx"
    # Verify the full content is preserved (heading + paragraph text)
    assert "Quy trinh su co" in chunks[0].text
    assert "Bao cao trong 24 gio." in chunks[0].text


def test_empty_markdown_is_failed_not_needs_ocr(tmp_path: Path) -> None:
    service = IngestionService(
        DocumentRegistry(tmp_path / "registry.json"),
        MemoryChunkStore(),
    )

    document = service.ingest_bytes("empty.md", b"")

    assert document.status is DocumentStatus.FAILED


def test_force_reindex_replaces_chunks_without_creating_new_version(tmp_path: Path) -> None:
    store = MemoryChunkStore()
    service = IngestionService(DocumentRegistry(tmp_path / "registry.json"), store)
    content = b"Current policy"
    document = service.ingest_bytes("policy.md", content)
    store.all_chunks = []

    reindexed = service.ingest_bytes("policy.md", content, force=True)

    assert reindexed.version == document.version
    assert [chunk.text for chunk in store.all_chunks] == ["Current policy"]


def test_unprocessable_new_version_keeps_last_ready_chunks(
    tmp_path: Path, monkeypatch
) -> None:
    store = MemoryChunkStore()
    service = IngestionService(DocumentRegistry(tmp_path / "registry.json"), store)
    responses = iter([("Active policy", "application/pdf"), ("", "application/pdf")])
    monkeypatch.setattr(
        "ingestion.service.parse_document",
        lambda filename, content: next(responses),
    )

    active = service.ingest_bytes("policy.pdf", b"version-1")
    pending = service.ingest_bytes("policy.pdf", b"version-2-scan")

    assert active.status is DocumentStatus.READY
    assert pending.status is DocumentStatus.NEEDS_OCR
    assert [chunk.text for chunk in store.all_chunks] == ["Active policy"]


def test_upload_storage_failure_does_not_publish_document(
    tmp_path: Path, monkeypatch
) -> None:
    store = MemoryChunkStore()
    registry = DocumentRegistry(tmp_path / "registry.json")
    service = IngestionService(registry, store, tmp_path / "uploads")

    def fail_write(self: Path, content: bytes) -> int:
        raise OSError("disk full")

    monkeypatch.setattr(Path, "write_bytes", fail_write)

    with pytest.raises(OSError, match="disk full"):
        service.ingest_bytes("policy.md", b"New policy")

    assert registry.list() == []
    assert store.all_chunks == []


def test_ingestion_service_allows_reingestion_of_same_source(tmp_path: Path) -> None:
    """In single-user mode, re-ingesting the same source always succeeds."""
    store = MemoryChunkStore()
    service = IngestionService(DocumentRegistry(tmp_path / "registry.json"), store)
    first = service.ingest_bytes("policy.md", b"HR policy")

    second = service.ingest_bytes("policy.md", b"Updated policy content")

    assert first.id == second.id
    assert second.version == first.version + 1


def test_ingestion_persists_canonical_legal_coordinates(tmp_path: Path) -> None:
    def build(document) -> None:
        document.add_heading("Chương I", level=1)
        document.add_heading("Điều 1. Phạm vi điều chỉnh", level=3)
        document.add_paragraph("Nội dung điều một.")

    store = MemoryChunkStore()
    service = IngestionService(DocumentRegistry(tmp_path / "registry.json"), store)

    document = service.ingest_bytes(
        "01_2021_ND-CP_283247.docx",
        _docx_bytes(build),
        {"canonical_doc_id": "01_2021_ND-CP_283247.md"},
    )
    chunks = store.list_document_chunks(document.id, document.version)

    assert chunks
    assert {chunk.coordinates.doc_id for chunk in chunks} == {"01_2021_ND-CP_283247.md"}
    article_chunks = [chunk for chunk in chunks if chunk.coordinates.article is not None]
    assert article_chunks
    assert all(chunk.coordinates.chapter is not None for chunk in article_chunks)


def test_parent_child_chunking_config_flows_through_ingestion(tmp_path: Path) -> None:
    store = MemoryChunkStore()
    service = IngestionService(
        DocumentRegistry(tmp_path / "registry.json"),
        store,
        chunking_config=ChunkingConfig(max_chars=20, parent_child_enabled=True, parent_max_chars=6000),
    )

    service.ingest_bytes("policy.md", b"Noi dung chinh sach nghi phep rat dai can duoc cat nho ra.")

    assert store.all_chunks
    assert all(chunk.parent_text is not None for chunk in store.all_chunks)


def test_raptor_disabled_by_default_adds_no_summary_nodes(tmp_path: Path) -> None:
    store = MemoryChunkStore()
    service = IngestionService(DocumentRegistry(tmp_path / "registry.json"), store)

    service.ingest_bytes(
        "policy.md",
        b"# A\nmot.\n# B\nhai.\n# C\nba.\n# D\nbon.\n# E\nnam.\n# F\nsau.",
    )

    assert not any(chunk.section and chunk.section.startswith("__raptor_summary_L") for chunk in store.all_chunks)


def test_raptor_enabled_appends_summary_nodes(tmp_path: Path) -> None:
    store = MemoryChunkStore()
    service = IngestionService(
        DocumentRegistry(tmp_path / "registry.json"),
        store,
        raptor_config=RaptorConfig(enabled=True, cluster_size=2, max_depth=1, provider=StubGenerationProvider()),
    )

    document = service.ingest_bytes(
        "policy.md",
        b"# A\nmot.\n# B\nhai.\n# C\nba.\n# D\nbon.\n# E\nnam.\n# F\nsau.",
    )

    summary_nodes = [
        chunk for chunk in store.all_chunks if chunk.section == "__raptor_summary_L1__"
    ]
    assert summary_nodes
    assert all(node.document_id == document.id for node in summary_nodes)
